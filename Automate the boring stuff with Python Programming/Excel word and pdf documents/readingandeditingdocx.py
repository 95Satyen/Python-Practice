# -*- coding: utf-8 -*-
"""
Created on Fri Feb 14 00:39:26 2020

@author: satye
"""

import docx
d = docx.Document('C:\\Users\\satye.MSI\\Documents\\GitHub\\Resume is LIT\\Resume  baX.docx')
#print(d.paragraphs)#list of paragraph objects
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[0].bold==True)

p.runs[1].underline=True
#d.save('C:\\Users\\satye.MSI\\Documents\\resumeupdate.docx')

d2 = docx.Document() #create new document object
d2.add_paragraph('hi')
d2.add_paragraph('this is it!')
d2.save('C:\\Users\\satye.MSI\\Documents\\yo.docx')
p = d2.paragraphs[0]
p.add_run('This is a new run.')#appending after something
print(p.runs)
p.runs[1].bold = True #set first paragraph as bold
d2.save('C:\\Users\\satye.MSI\\Documents\\yo.docx')


################
#copy text
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        return '\n'.join(fullText)
print(getText('C:\\Users\\satye.MSI\\Documents\\yo.docx'))    
